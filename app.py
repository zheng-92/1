import os
import sys
import sqlite3
import streamlit as st
from datetime import datetime
from typing import Annotated, List, TypedDict, Literal
from operator import add
import re
import time
from concurrent.futures import ThreadPoolExecutor
from docx import Document as DocxDocument
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'
os.environ['HF_HOME'] = r'D:\zjq\python\RAG\models_cache'

from pydantic import BaseModel, Field
from langgraph.graph import StateGraph, END
from langgraph.types import Send
from langchain_core.output_parsers import PydanticOutputParser
from langgraph.checkpoint.sqlite import SqliteSaver
from langchain_core.messages import BaseMessage, HumanMessage, AIMessage, SystemMessage
from langchain_openai import ChatOpenAI
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_community.retrievers import BM25Retriever
from langchain_community.document_loaders import TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sentence_transformers import CrossEncoder
from dotenv import load_dotenv
from langchain_core.documents import Document
from docx.shared import Pt
from docx.oxml.ns import qn

load_dotenv()

SUMMARIZER_PROMPT = """# ROLE: 法律案件备忘录管理员
# TASK: 维护一份精简的长期记忆。
# 备忘录格式：
1. 核心事实：[人物、公司名、持股比例、涉及金额等]
2. 矛盾点：[目前最主要的法律争议]
3. 已知诉求：[用户明确表达过想要什么，如：想要工资、不想起草函件]
4. 对话阶段：[已经聊到了哪个阶段，避免重复确认]

# RULES:
- 每次根据新对话更新内容，删除过时的推测。
- 严禁分析，只记录事实。"""

ROUTER_PROMPT = """
# ROLE: 专业对话分类师
# TASK: 判定用户意图。
1. 如果是关于《公司法》、股权、公司治理、企业注销等，回答：LEGAL。
2. 如果是关于离婚、继承、侵权、刑事犯罪等非公司法领域，回答：OUT_OF_SCOPE。
3. 如果是普通聊天、打招呼，回答：CHAT。
只回答一个标签，不要解释。
"""

REWRITER_PROMPT = """# ROLE: 法律语境补全专家
# TASK: 将模糊口语重写为独立的专业查询语句。根据备忘录补全指代（他/我公司等）。"""

PLANNER_PROMPT = """# ROLE: 资深律师工作流规划师
# TASK: 将法律需求拆解为 2 个精准子任务。
#  严格约束：
1. 你必须假设世界上只有《中华人民共和国公司法》及其司法解释。
2. 严禁在任务描述中提及《登记管理条例》、《民法典》或其他任何法律。
3. 拆解的任务必须聚焦于：从现有库中检索对应的公司法条文。
# FORMAT:
1. [evidence] 负责查找《公司法》中关于[具体问题]的法律依据。
2. [strategy] 负责基于公司法给出建议或框架。"""


EVIDENCE_AGENT_PROMPT = """# ROLE: 严谨的法学研究员
# TASK: 提取法律依据。
# 绝对禁令：
1. 严禁引用任何非《公司法》及其司法解释的法律（如：禁止引用《登记管理条例》、《民法典》、《刑法》等）。
2. 如果用户问题在《公司法》中找不到依据，请直接回答：[未找到相关法条]。
3. 严禁凭空想象具体的条款号，必须只根据公司法框架提出检索关键词。
4.严禁凭空想象或输出任何具体的法条序号（如“根据第七十八条”），因为你目前没有挂载知识库。
5. 你的输出只能是检索方向，例如：“为了解决这个问题，需要重点查证公司法中关于XX的规定”。"""

STRATEGY_AGENT_PROMPT = """# ROLE: 实战派法律顾问
# TASK: 给出可落地的行动方案。
# 限制规则：
1. 仅限依据《公司法》提供建议。
2. 严禁引用你自身知识库中关于行政登记管理、劳动、侵权等领域的外部条例。
3. 如果任务超纲，必须输出：[超纲任务：此问题属于非公司法领域]。
4.严禁引用任何具体的法条序号！只提供业务操作步骤"""

EXECUTOR_PROMPT = """# ROLE: 资深法律顾问（温暖且专业）
# CONTEXT: 你在即时通讯软件上提供咨询，用户可能焦虑或困惑。

# GOALS:
1. 像真人一样接话：如果用户表达了情绪或拒绝（如“不需要”），先给出 1 句共情回应，再继续。
2. 拒绝复读机：严禁在每句话开头说“我已经记住了...”。
3. 知识与交互平衡：
   - 对于 CHAT：像朋友一样自然闲聊，利用备忘录确认关键点，顺势引导。
   - 对于 LEGAL：严谨分析，使用“### ⚖️ 法律分析”标签。必须严格依据下方的调研报告回答，答案中的每一个法律推导都必须在报告中找到依据。如果报告中没有相关法条，请明确说明“法条库未涵盖”，严禁脱离报告自由发挥,必须按知识库里有的输出，严禁使用你自己的内部知识去输出。
4. 对于 LEGAL 咨询：必须严格依据调研报告中的【参考依据】回答！答案中的每一个法律推导都必须在【参考依据】中找到原文。如果【参考依据】中没有相关法条，请明确说明“法条库未涵盖”，严禁脱离依据自由发挥！如果【查证专家结论】与【参考依据】冲突，必须绝对以【参考依据】为准！
5. 拦截红线：如果报告中包含“[未找到相关法条]”或“[超纲任务]”，你必须回复：“抱歉，我的知识库目前仅涵盖《公司法》及相关领域，您提及的问题（如侵权纠纷等）超出了我的专业范围，无法为您解答。” 严禁自由发挥！    

# 文书起草 (SKILL)：
- 在用户明确指令时，在 doc_content 中生成文书
- 文书正文包裹在 <DOC_CONTENT> 和 </DOC_CONTENT> 内，纯文本公文格式。
- 最后一行添加标记：[TRIGGER_DOC: 文书名称]。
- 写入文书时，不要带自己的思考结果，只留符合文书的正确内容。
- 禁止在 reply_text 中展示文书的任何具体正文内容
- 绝对禁止在 reply_text 或 doc_content 中输出任何 HTML/XML 标签（如 <DOC_CONTENT>）
"""



class AgentState(TypedDict):
    messages: Annotated[List[BaseMessage], add]
    context: List[str]
    intent: str
    rewritten_query: str
    sub_tasks: List[str]
    evidence_results: Annotated[List[str], add]
    strategy_results: Annotated[List[str], add]
    summary: str
class FinalResponse(BaseModel):
    reply_text: str = Field(description="给用户的聊天回复正文，包含共情和法律分析。")
    needs_doc: bool = Field(description="用户是否明确要求起草、生成或写一份文书？")
    doc_title: str = Field(description="如果要生成文书，提取文书的标题，例如'解除劳动合同通知书'。如果不生成则为空。")
    doc_content: str = Field(description="如果要生成文书，这里是纯文本正文。必须100%忠实于用户提供的金额、姓名、日期等核心事实。绝对不要使用任何Markdown符号（如**或#等）。")
def init_db():
    conn = sqlite3.connect("law_agent_memory.db", check_same_thread=False)
    conn.execute(
        "CREATE TABLE IF NOT EXISTS chat_sessions (thread_id TEXT PRIMARY KEY, title TEXT, updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP)")
    conn.commit()
    return conn


db_conn = init_db()


@st.cache_resource
def init_core_components():
    embeddings = HuggingFaceEmbeddings(model_name="BAAI/bge-large-zh-v1.5", model_kwargs={'device': 'cpu'})
    vectorstore = Chroma(persist_directory="D:/zjq/python/RAG/law_chroma_db", embedding_function=embeddings)

    # === BM25 直接读取 Chroma 里的结构化数据 ===
    db_data = vectorstore.get()
    all_docs = []
    if db_data and 'documents' in db_data and db_data['documents']:
        for i in range(len(db_data['documents'])):
            all_docs.append(Document(page_content=db_data['documents'][i], metadata=db_data['metadatas'][i]))
    else:
        print(" 警告：向量库为空，请先运行 build_vector_db.py")

    bm25 = BM25Retriever.from_documents(all_docs)

    # Reranker 模型初始化
    reranker = CrossEncoder('BAAI/bge-reranker-v2-m3', device='cpu')

    api_key = os.getenv("DEEPSEEK_API_KEY")
    llm = ChatOpenAI(model="deepseek-chat", api_key=api_key, base_url="https://api.deepseek.com", temperature=0)
    memory = SqliteSaver(sqlite3.connect("law_agent_memory.db", check_same_thread=False))

    return vectorstore.as_retriever(search_kwargs={"k": 10}), bm25, llm, memory, reranker


vector_retriever, bm25_retriever, llm, memory, reranker = init_core_components()


def summarize_node(state: AgentState):
    messages = state["messages"]
    existing_summary = state.get("summary", "")
    # 1. 如果是第一轮对话，不需要总结
    if len(messages) <= 2 and not existing_summary:
        return {"summary": "暂无备忘记录"}

    # 2. 每隔 2 轮对话（即新增 4 条消息：2条user，2条assistant），才触发一次总结
    if len(messages) % 4 != 0 and existing_summary:
        return {"summary": existing_summary}  # 直接复用旧的备忘录，不调用大模型

    # 触发总结：只把最近的 4 条消息拿去总结，增量更新
    recent_messages = messages[-4:]
    res = llm.invoke([SystemMessage(content=SUMMARIZER_PROMPT),
                      HumanMessage(content=f"已有备忘：{existing_summary}\n新对话：{recent_messages}")]).content
    return {"summary": res}


def router_node(state: AgentState):
    query = state["messages"][-1].content
    res = llm.invoke([SystemMessage(content=ROUTER_PROMPT), HumanMessage(content=query)]).content.strip().upper()
    if "OUT_OF_SCOPE" in res:
        intent = "OUT_OF_SCOPE"
    elif "LEGAL" in res:
        intent = "LEGAL"
    else:
        intent = "CHAT"

    print(f"\n--- 意图决策: {intent} ---")
    return {"intent": intent}


def rewriter_node(state: AgentState):
    if state["intent"] in ["CHAT", "OUT_OF_SCOPE"]:
        return {"rewritten_query": ""}

    res = llm.invoke([SystemMessage(content=f"背景：{state.get('summary', '')}\n{REWRITER_PROMPT}"),
                      HumanMessage(content=state["messages"][-1].content)]).content
    return {"rewritten_query": res}


def planner_node(state: AgentState):
    if state["intent"] in ["CHAT", "OUT_OF_SCOPE"]:
        return {"sub_tasks": []}

    res = llm.invoke([SystemMessage(content=PLANNER_PROMPT),
                      HumanMessage(content=f"背景：{state['summary']}\n问题：{state['rewritten_query']}")]).content
    return {"sub_tasks": [l for l in res.split('\n') if '[' in l]}


#  并发协作
def evidence_node(state: AgentState):
    task = state["sub_tasks"][0] # 传递进来的具体任务
    print(f"\n[Evidence Agent]  正在查证法条: {task}")
    content = llm.invoke([SystemMessage(content=EVIDENCE_AGENT_PROMPT), HumanMessage(content=f"任务：{task}")]).content
    return {"evidence_results": [content]}

def strategy_node(state: AgentState):
    task = state["sub_tasks"][0]
    print(f"[Strategy Agent] 💡 正在制定策略: {task}")
    draft_keywords = ["起草", "生成", "写一份", "写", "拟定"]
    is_drafting = any(kw in task for kw in draft_keywords)
    p = STRATEGY_AGENT_PROMPT + ("" if is_drafting else "\n 仅提供建议，严禁输出公文模板，严禁使用 <DOC_CONTENT> 标签。")
    content = llm.invoke([SystemMessage(content=p), HumanMessage(content=f"背景：{state.get('summary','')}\n任务：{task}")]).content
    return {"strategy_results": [content]}


def map_tasks(state: AgentState):
    """
    动态路由逻辑：根据 Planner 拆分的子任务，并行分发给对应的 Agent Node
    """
    if state["intent"] in ["CHAT", "OUT_OF_SCOPE"] or not state.get("sub_tasks"):
        return "executor"

    send_actions = []
    print(f"[Graph] 🚀 正在并发分发 {len(state['sub_tasks'])} 个子任务")
    for task in state["sub_tasks"]:
        if "[evidence]" in task:
            send_actions.append(Send("evidence", {"sub_tasks": [task], "summary": state.get("summary", "")}))
        elif "[strategy]" in task:
            send_actions.append(Send("strategy", {"sub_tasks": [task], "summary": state.get("summary", "")}))

    return send_actions if send_actions else "executor"


def executor_node(state: AgentState):
    query = state["messages"][-1].content
    short_term_context = state["messages"][-4:-1]

    if state["intent"] == "OUT_OF_SCOPE":
        refuse_msg = "抱歉，我的知识库目前仅涵盖《公司法》及相关领域，您提及的问题超出了我的专业范围，无法为您解答。"
        print("[Executor]  检测到非公司法领域的超纲问题，已直接拒绝。")
        return {"messages": [AIMessage(content=refuse_msg)]}

    parser = PydanticOutputParser(pydantic_object=FinalResponse)
    format_instructions = parser.get_format_instructions()

    # 强制 JSON 指令
    system_msg = SystemMessage(
        content=f"{EXECUTOR_PROMPT}\n\n当前客户备忘录：\n{state.get('summary', '尚无记录')}\n\n"
                f"【绝对指令】你必须且只能输出严格的 JSON 字符串对象，请直接以 {{ 开头，以 }} 结尾。禁止使用 Markdown 标记（如 ```json）。\n"
                f"格式说明：\n{format_instructions}"
    )

    json_llm = llm.bind(response_format={"type": "json_object"})

    if state["intent"] == "CHAT":
        prompt = [system_msg] + short_term_context + [HumanMessage(content=query)]
        ans = json_llm.invoke(prompt)
        try:
            parsed_ans = parser.invoke(ans.content)  # 注意这里改成了 ans.content
            return {"messages": [AIMessage(content=parsed_ans.reply_text)]}
        except Exception as e:
            print(f"CHAT解析失败: {e}")
            return {"messages": [ans]}

    # LEGAL 模式
    v_docs = vector_retriever.invoke(state["rewritten_query"])
    b_docs = bm25_retriever.invoke(state["rewritten_query"])
    candidate_docs = {d.page_content: d for d in (v_docs + b_docs)}.values()
    top_docs = []
    if candidate_docs:
        doc_list = list(candidate_docs)
        pairs = [[state["rewritten_query"], d.page_content] for d in doc_list]
        scores = reranker.predict(pairs)
        ranked_results = sorted(zip(scores, doc_list), key=lambda x: x[0], reverse=True)
        top_docs = [res[1] for res in ranked_results[:3]]

    context_text = "\n".join([d.page_content for d in top_docs])
    e_res = state.get('evidence_results', ['无'])
    s_res = state.get('strategy_results', ['无'])
    final_report = f"【参考依据】：{context_text}\n【查证专家结论】：{e_res}\n【策略专家建议】：{s_res}"

    prompt = [system_msg] + short_term_context + [
        HumanMessage(content=f"请严格基于以下调研报告的【参考依据】回答用户问题：{query}\n\n调研报告：\n{final_report}")
    ]

    ans = json_llm.invoke(prompt)
    try:
        parsed_ans = parser.invoke(ans.content)

        e_res_str = "".join(state.get('evidence_results', []))
        s_res_str = "".join(state.get('strategy_results', []))

        if "[超纲任务]" in s_res_str or "[未找到相关法条]" in e_res_str:
            refuse_msg = "抱歉，知识库目前仅涵盖《公司法》及其司法解释，超出了专业范围。"
            return {"messages": [AIMessage(content=refuse_msg)]}

        final_text = parsed_ans.reply_text
        if parsed_ans.needs_doc and parsed_ans.doc_content:
            final_text += f"\n\n<DOC_CONTENT>\n{parsed_ans.doc_content.strip()}\n</DOC_CONTENT>\n[TRIGGER_DOC: {parsed_ans.doc_title.strip()}]"

        print("[Executor] 响应构建完成 (JSON解析成功)")
        return {"messages": [AIMessage(content=final_text)], "context": [d.page_content for d in top_docs]}
    except Exception as e:
        print(f"[Executor]  兜底降级 (异常: {e})")
        return {"messages": [ans], "context": [d.page_content for d in top_docs] if top_docs else []}


# --- 图构建逻辑  ---
workflow = StateGraph(AgentState)
workflow.add_node("summarize", summarize_node)
workflow.add_node("router", router_node)
workflow.add_node("rewriter", rewriter_node)
workflow.add_node("planner", planner_node)
# 添加拆分后的两个独立 Agent 节点
workflow.add_node("evidence", evidence_node)
workflow.add_node("strategy", strategy_node)
workflow.add_node("executor", executor_node)

workflow.set_entry_point("summarize")
workflow.add_edge("summarize", "router")
workflow.add_edge("router", "rewriter")
workflow.add_edge("rewriter", "planner")

# 核心修改：使用 Conditional Edges 实现动态并行分发
workflow.add_conditional_edges("planner", map_tasks, ["evidence", "strategy", "executor"])

# 两个并行节点执行完毕后，统一汇聚到 executor
workflow.add_edge("evidence", "executor")
workflow.add_edge("strategy", "executor")
workflow.add_edge("executor", END)

app = workflow.compile(checkpointer=memory)



def apply_fixed_style():
    st.markdown("""<style>
        [data-testid="stAppViewContainer"] { background-color: #0E1117; }
        .welcome-title { font-size: 42px !important; font-weight: 800 !important; background: linear-gradient(90deg, #58A6FF, #BC8CF2); -webkit-background-clip: text; -webkit-text-fill-color: transparent; }
        .stChatMessage { background-color: #161B22 !important; border: 1px solid #30363D !important; border-radius: 10px !important; }
    </style>""", unsafe_allow_html=True)


def stream_display(text):
    placeholder = st.empty();
    full_res = ""
    for chunk in re.split(r'(\s+)', text):
        full_res += chunk;
        placeholder.markdown(full_res + "▌");
        time.sleep(0.01)
    placeholder.markdown(full_res)


def create_legal_word(content: str, doc_type: str):
    doc = DocxDocument()
    # 1. 彻底提取并清洗正文
    match = re.search(r'<DOC_CONTENT>(.*?)</DOC_CONTENT>', content, re.DOTALL)
    text = match.group(1).strip() if match else content
    lines = text.split('\n')
    if lines and (doc_type in lines[0] or "标题" in lines[0]):
        text = '\n'.join(lines[1:]).strip()
    text = re.sub(r'</?DOC_CONTENT>', '', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'#+\s*(.*?)\n', r'\1\n', text)
    text = re.sub(r'\[TRIGGER_DOC:.*?\]', '', text)
    # 2. 深度清洗标题 (doc_type)
    clean_title = "".join(re.findall(r'[\u4e00-\u9fa5]+', doc_type))
    if not clean_title: clean_title = "法律文书"
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title_paragraph.add_run(clean_title)
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.name = 'SimHei'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    # 4. 设置正文样式
    paragraphs = text.split('\n')
    for p_text in paragraphs:
        if not p_text.strip(): continue
        p = doc.add_paragraph()
        run = p.add_run(p_text.strip())
        run.font.size = Pt(12)
        run.font.name = 'SimSun'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # 保存文档
    path = f"Legal_Draft_{datetime.now().strftime('%H%M%S')}.docx"
    doc.save(path)
    return path

def get_all_sessions():
    try:
        cursor = db_conn.cursor()
        cursor.execute("SELECT thread_id, title FROM chat_sessions ORDER BY updated_at DESC")
        return [{"id": row[0], "title": row[1]} for row in cursor.fetchall()]
    except:
        return []


def save_title(thread_id, query):
    title = llm.invoke([HumanMessage(content=f"总结6字内标题：{query}")]).content.strip()
    db_conn.execute("INSERT OR REPLACE INTO chat_sessions (thread_id, title) VALUES (?, ?)", (thread_id, title))
    db_conn.commit()


def main():
    st.set_page_config(page_title="LawAgent Pro", page_icon="⚖️", layout="wide")
    apply_fixed_style()

    if "messages" not in st.session_state: st.session_state.messages = []
    if "current_thread" not in st.session_state: st.session_state.current_thread = f"T_{int(time.time())}"

    with st.sidebar:
        st.title("🏛️ LawAgent Pro")
        if st.button("➕ 开启新咨询", use_container_width=True):
            st.session_state.current_thread = f"T_{int(time.time())}"
            st.session_state.messages = [];
            st.rerun()
        st.markdown("---")
        sessions = get_all_sessions()
        session_ids = [s["id"] for s in sessions]
        if st.session_state.current_thread not in session_ids:
            session_ids.insert(0, st.session_state.current_thread)
            sessions.insert(0, {"id": st.session_state.current_thread, "title": "✨ 当前对话"})

        selected_id = st.selectbox("历史记录", options=session_ids,
                                   index=session_ids.index(st.session_state.current_thread),
                                   format_func=lambda x: {s["id"]: s["title"] for s in sessions}.get(x, x))

        if selected_id != st.session_state.current_thread:
            st.session_state.current_thread = selected_id;
            st.session_state.messages = []
            state = app.get_state({"configurable": {"thread_id": selected_id}})
            if state and "messages" in state.values:
                for m in state.values["messages"]:
                    st.session_state.messages.append(
                        {"role": "user" if isinstance(m, HumanMessage) else "assistant", "content": m.content})
            st.rerun()

    if not st.session_state.messages:
        st.markdown(
            '<div style="text-align:center; padding:50px;"><div class="welcome-title">🏛️ LawAgent Pro</div><p style="color:#8B949E">基于 Rerank + 并发架构的专家级法律助手</p></div>',
            unsafe_allow_html=True)

    for msg in st.session_state.messages:
        with st.chat_message(msg["role"]):
            disp = re.sub(r'<DOC_CONTENT>.*?</DOC_CONTENT>', '\n\n*(📝 详细文书已生成)*\n\n', msg["content"],
                          flags=re.DOTALL)
            st.markdown(re.sub(r'\[TRIGGER_DOC:.*?\]', '', disp))

    if user_input := st.chat_input("请输入咨询需求..."):
        st.session_state.messages.append({"role": "user", "content": user_input})
        with st.chat_message("user"):
            st.markdown(user_input)

        with st.chat_message("assistant"):
            with st.status("⚡ 收到需求，正在启动系统...", expanded=False) as status:
                if len(st.session_state.messages) == 1:
                    status.update(label="📝 正在为本次咨询生成档案标题...")
                    save_title(st.session_state.current_thread, user_input)

                status.update(label="🔍 专家组协作分析中...")
                config = {"configurable": {"thread_id": st.session_state.current_thread}}
                final_state = app.invoke({"messages": [HumanMessage(content=user_input)]}, config=config)
                raw_ans = final_state["messages"][-1].content
                status.update(label="✅ 分析完成", state="complete")

            display_text = re.sub(r'<DOC_CONTENT>.*?</DOC_CONTENT>', '\n\n> 📝 **文书已就绪**，请点击下方按钮下载。',
                                  raw_ans, flags=re.DOTALL)
            display_text = re.sub(r'\[TRIGGER_DOC:.*?\]', '', display_text)
            stream_display(display_text)
            st.session_state.messages.append({"role": "assistant", "content": raw_ans})

            if "[TRIGGER_DOC:" in raw_ans:
                doc_name = raw_ans.split("[TRIGGER_DOC:")[1].split("]")[0]
                path = create_legal_word(raw_ans, doc_name)
                with open(path, "rb") as f:
                    st.download_button(f"📥 下载《{doc_name}》", f, file_name=path, use_container_width=True)

if __name__ == "__main__":
    main()